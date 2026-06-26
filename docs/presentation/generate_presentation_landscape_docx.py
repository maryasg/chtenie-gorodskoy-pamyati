#!/usr/bin/env python3
"""Landscape Word: slide image left, speech text right (slides 13-17)."""

from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
PDF = ROOT / "pitch" / "!Презентация_Читать_стены.pdf"
IMG_DIR = ROOT / "slide_images"
OUTPUT = ROOT / "Речь_презентация_Мария_горизонталь.docx"

# PDF page index (0-based) -> slide number on deck
SLIDES = [
    (13, 13, "Технологическая составляющая · MVP"),
    (14, 14, "Archiview"),
    (15, 15, "Будущее проекта"),
    (16, 16, "Прототип · путь пользователя (1)"),
    (17, 17, "Прототип · путь пользователя (2)"),
]

SPEECH = {
    13: [
        "Меня зовут Мария, я отвечала за технологическую реализацию и цифровой прототип.",
        "На слайде — не концепт, а рабочий MVP: веб-сайт и десктопное приложение Archiview.",
        "Задача технологии — дать эксперту инструмент размечать на фасаде следы памяти, "
        "а пользователю — увидеть их на сайте и в перспективе в дополненной реальности.",
    ],
    14: [
        "Archiview мы создали специально для проекта. Это инструмент разметки фасада на базе OpenCV.",
        "Эксперт выбирает пару снимков, готовит и выпрямляет их, сравнивает исторический "
        "и современный слой, обводит полигоны следов памяти — с типом и интерпретацией каждого элемента.",
        "Результат экспортируется на сайт. Так мы собираем базу размеченных фасадов — "
        "основу и для публичного прототипа, и для будущего развития.",
        "Идите по схеме на слайде: выбор фото → подготовка → сравнение → разметка → экспорт.",
    ],
    15: [
        "Сейчас в MVP разметка ручная — и это фундамент: мы накапливаем датасет и словарь визуальных признаков.",
        "В перспективе — computer vision, в том числе обучение на Roboflow: помогать находить следы "
        "на зданиях, которых ещё нет в базе — как гипотезы для эксперта, не как автоматический вердикт.",
        "По AR: сейчас на сайте AR-preview — симуляция подсветки на полевом фото. "
        "Дальше — настоящая AR по геолокации для зданий из базы.",
        "Отдельно — доступность и инклюзия в ближайших планах.",
    ],
    16: [
        "На сайте пользователь начинает с карты четырёх пилотных зданий, переходит в карточку здания, "
        "видит интерактивный фасад с подсветкой размеченных зон и открывает карточку следа.",
        "В карточке каждого элемента есть блок «Достоверность» — например, «Подтверждено». "
        "Мы заранее договорились, насколько надёжны источники: акт экспертизы, архивный снимок "
        "или пока гипотеза. Пользователь сразу видит, чему можно доверять.",
    ],
    17: [
        "Дальше — сравнение до и после: видно, что изменилось на фасаде.",
        "Блок слоёв времени — архивные фото и планы сквозь годы.",
        "Текстовые исторические слои — интерпретация и контекст.",
        "И источники — пользователь может сам сверить материалы и сделать вывод. "
        "Сайт не заменяет эксперта, но показывает, на чём держится каждый вывод.",
    ],
}


def export_slide_images() -> None:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(PDF)
    for page_idx, slide_num, _ in SLIDES:
        page = doc[page_idx]
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        pix.save(str(IMG_DIR / f"slide_{slide_num:02d}.png"))


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)


def add_speech_cell(cell, paragraphs: list[str]) -> None:
    cell.text = ""
    for i, para in enumerate(paragraphs):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.text = para
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.size = Pt(11)


def main() -> None:
    export_slide_images()

    doc = Document()
    set_landscape(doc.sections[0])

    title = doc.add_heading("Речь Maria · слайды 13–17", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Формат: слайд слева · текст справа · альбомная ориентация")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for i, (page_idx, slide_num, subtitle) in enumerate(SLIDES):
        if i > 0:
            doc.add_page_break()

        heading = doc.add_heading(f"Слайд {slide_num}. {subtitle}", level=1)

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        left, right = table.rows[0].cells
        left.width = Cm(14.5)
        right.width = Cm(13.5)

        img_path = IMG_DIR / f"slide_{slide_num:02d}.png"
        if img_path.exists():
            left.paragraphs[0].add_run().add_picture(str(img_path), width=Cm(14))
        else:
            left.text = "(изображение слайда не найдено)"

        label = right.paragraphs[0]
        label_run = label.add_run("Речь")
        label_run.bold = True
        label_run.font.size = Pt(12)

        add_speech_cell(right, SPEECH[slide_num])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
